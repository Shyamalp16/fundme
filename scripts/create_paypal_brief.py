from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)

INK = RGBColor(23, 23, 20)
MUTED = RGBColor(93, 91, 84)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
CORAL = RGBColor(199, 53, 37)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_paragraph(paragraph, before=0, after=6, line=1.1, alignment=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def add_text(doc, text, before=0, after=6, size=11, color=INK, bold=False, italic=False):
    p = doc.add_paragraph()
    style_paragraph(p, before=before, after=after)
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich(doc, parts, before=0, after=6):
    p = doc.add_paragraph()
    style_paragraph(p, before=before, after=after)
    for text, opts in parts:
        r = p.add_run(text)
        set_font(r, **opts)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    if level == 1:
        size, color, before, after = 16, BLUE, 16, 8
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, DARK_BLUE, 8, 4
    style_paragraph(p, before=before, after=after)
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=True)
    return p


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.05)
        r = p.add_run(header)
        set_font(r, size=9.5, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            style_paragraph(p, after=0, line=1.08)
            r = p.add_run(value)
            set_font(r, size=9.5, color=INK)
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=100, bottom=100)
    return table


def set_header_footer(section):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(hp, after=0)
    r = hp.add_run("fudnME  |  PAYPAL PLATFORM APPROVAL BRIEF")
    set_font(r, size=8, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(fp, after=0)
    r = fp.add_run("Pre-launch / confidential working brief  |  August 16, 2026")
    set_font(r, size=8, color=MUTED)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_header_footer(section)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1


def build_docx():
    doc = Document()
    configure_styles(doc)

    add_text(doc, "PAYPAL PLATFORM APPROVAL BRIEF", before=10, after=5, size=9, color=CORAL, bold=True)
    title = doc.add_paragraph()
    style_paragraph(title, after=5, line=1.0)
    r = title.add_run("fudnME")
    set_font(r, name="Arial", size=28, color=INK, bold=True)
    r = title.add_run("  |  Canadian pre-launch micro-gifting platform")
    set_font(r, size=15, color=MUTED)
    add_text(doc, "Request for preliminary platform, compliance, and product-fit review", after=14, size=12, color=MUTED, italic=True)

    add_table(
        doc,
        ["Field", "Current information"],
        [
            ("Applicant", "fudnME — pre-incorporation Canadian project"),
            ("Status", "Pre-launch; no live payments accepted; prototype deployed"),
            ("Website", "https://fudnme.now"),
            ("Launch market", "Canada-first; proposed CAD-only beta; adults 18+"),
            ("Requested PayPal path", "PayPal Complete Payments / multiparty platform review"),
            ("Founder contact", "[Insert founder legal name, business email, phone]"),
        ],
        [2700, 6660],
    )

    add_heading(doc, "1. Executive request", 1)
    add_text(doc, "fudnME is seeking a preliminary written determination from PayPal on whether its proposed Canadian web platform can be onboarded as a permitted platform or marketplace using PayPal Complete Payments / multiparty capabilities. We are not asking to activate live payments in this review. We want to understand the correct PayPal product, classification, approval path, and required risk controls before incorporating and building production payment infrastructure.")
    add_text(doc, "Please route this inquiry to the team responsible for Platforms and Marketplaces, crowdfunding or personal-fundraising review, risk/compliance, and Canadian onboarding. We would appreciate a case or approval reference number and a written response to the questions in Section 8.", after=8)

    add_heading(doc, "2. Product summary", 1)
    add_text(doc, "fudnME is a social micro-gifting platform. An adult creator posts a public, lighthearted personal goal and an aspirational target. Other adults may voluntarily give a small, unconditional monetary gift to that ask. The core mechanic is intentionally limited: a supporter may contribute a maximum of CAD $1 to a particular ask and may support up to ten asks per week.")
    add_text(doc, "The product is designed to be social, shareable, and low-pressure. Applications may be funny, random, genuine, or aspirational. The platform does not promise that a target will be reached and does not verify how every gift is ultimately spent.")
    add_text(doc, "The deployed prototype is a pre-launch explanation and waitlist site only. It currently accepts no money and creates no payment account. Payments would be enabled only after provider approval, recipient onboarding, identity checks, moderation, fraud controls, payout rules, and legal review are complete.", after=8)

    add_heading(doc, "3. What a contribution is — and is not", 1)
    add_table(
        doc,
        ["Contribution treatment", "fudnME position"],
        [
            ("Nature", "Voluntary, unconditional personal gift from one user to another."),
            ("Supporter receives", "Nothing: no product, service, reward, perk, equity, ownership, interest, loan repayment, prize, raffle entry, or chance to win."),
            ("Charity status", "Not presented as a charitable donation and not represented as tax deductible."),
            ("Target", "Aspirational lifetime goal; failure to reach it would not automatically trigger a refund."),
            ("Discovery", "Feed ordering may use curated or randomized discovery, but randomness never determines who receives money or any benefit."),
            ("Prohibited content", "Illegal activity, minors, impersonation, harassment, hate, sexual exploitation, medical claims, tragedy exploitation, fraud, misleading claims, and privacy violations."),
        ],
        [2700, 6660],
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "4. Proposed V1 operating model", 1)
    add_table(
        doc,
        ["Area", "Proposed approach for provider review"],
        [
            ("Geography and currency", "Canada-first beta; one campaign currency, proposed CAD; no internal foreign-exchange wallet in V1."),
            ("Recipient onboarding", "Recipients onboarded through PayPal-supported partner/multiparty onboarding and any required KYC/KYB checks before payout."),
            ("Payment limit", "One contribution per supporter per ask, capped at CAD $1; up to ten supported asks per supporter per week."),
            ("Payout timing", "Proposed 7–14 day fraud/chargeback review, minimum payout threshold, and weekly payout schedule; exact values to be set with PayPal."),
            ("Funds movement", "Use PayPal platform APIs for payment capture, recipient allocation, disbursement, refunds, disputes, and reporting; fudnME does not maintain a stored-value wallet."),
            ("Platform revenue", "A separately disclosed platform fee is being evaluated. Preferred discussion model: the creator receives the stated CAD $1 gift amount and any platform fee is disclosed separately, subject to PayPal, legal, and app-store review."),
            ("Refunds and chargebacks", "Refunds for fraud, duplicate payments, technical errors, policy violations, or as otherwise required by PayPal; fudnME will follow PayPal’s reserve, dispute, and negative-balance requirements."),
        ],
        [2700, 6660],
    )

    add_heading(doc, "5. Illustrative fund flow", 1)
    add_table(
        doc,
        ["Step", "Flow"],
        [
            ("1", "Supporter selects a public ask and starts PayPal checkout."),
            ("2", "PayPal processes the contribution and applies any approved platform-fee configuration."),
            ("3", "The contribution is associated with the onboarded recipient/payee and the specific ask."),
            ("4", "PayPal applies fraud, compliance, dispute, reserve, and delayed-disbursement controls."),
            ("5", "After the applicable review period and payout conditions, PayPal disburses the recipient balance."),
            ("6", "fudnME receives reporting for reconciliation and applies its moderation, account, and support policies."),
        ],
        [900, 8460],
        header_fill=LIGHT_BLUE,
    )
    add_text(doc, "This is a proposed operating model for review, not a claim that PayPal has approved the flow. We would adapt the model to PayPal’s required account structure, API capabilities, reserves, settlement timing, and risk allocation.", before=4, after=8, size=10, color=MUTED, italic=True)

    add_heading(doc, "6. Trust, safety, and fraud controls", 1)
    add_text(doc, "The initial product will be web-first and adults-only. Before live payments, fudnME plans to implement: creator identity and payout verification; email and account-age checks; card, device, IP, velocity, and account-linking signals; self-funding and collusion detection; duplicate-account controls; manual review queues; creator impersonation protection; content reporting and appeals; and clear refund, chargeback, account-ban, and payout-hold disclosures.")
    add_text(doc, "Moderation will prohibit illegal activity, hate and harassment, doxxing, sexual exploitation, minors, impersonation, misleading or fabricated claims, medical fundraising claims, self-harm exploitation, copyright violations, and any use of the platform as a contest, raffle, prize, investment, loan, or commercial sale. Public submissions will be subject to removal and escalation procedures.")

    add_heading(doc, "7. Business and launch status", 1)
    add_table(
        doc,
        ["Item", "Status"],
        [
            ("Legal entity", "Not incorporated yet; contacting PayPal as a pre-incorporation project."),
            ("Production PayPal account", "Not opened; no personal account will be used to process production funds."),
            ("Prototype", "Deployed pre-launch website at https://fudnme.now; waitlist only; no real money accepted."),
            ("Transaction history", "None; the product is pre-revenue and has not processed live contributions."),
            ("Initial volume planning", "Illustrative planning range only: under CAD $100,000 gross contribution volume in the first 12 months; subject to validation and provider guidance."),
            ("Next legal step", "Consult Canadian payments/fintech counsel after initial provider feedback and before accepting live funds."),
        ],
        [2700, 6660],
        header_fill=LIGHT_GRAY,
    )

    add_heading(doc, "8. Questions for PayPal", 1)
    questions = [
        ("1", "How should PayPal classify this model in Canada: crowdfunding, personal fundraising, a marketplace/platform, a payment facilitator, or another category?"),
        ("2", "Is this model eligible for PayPal Complete Payments / multiparty capabilities in Canada, subject to approval?"),
        ("3", "Can PayPal onboard individual adult recipients/payees, perform the required KYC/AML checks, and disburse their balances?"),
        ("4", "Can PayPal support a CAD $1 contribution to a recipient, a separate disclosed platform fee, and the recipient receiving the full stated gift amount?"),
        ("5", "Can contributions be subject to delayed disbursement while fraud and chargeback review occurs? What holding period, reserve, and payout threshold would apply?"),
        ("6", "Who carries liability for refunds, disputes, chargebacks, negative balances, banned accounts, and unavailable recipients?"),
        ("7", "What countries, currencies, payment methods, and recipient locations are available for a Canada-first launch?"),
        ("8", "What content moderation, fraud monitoring, transaction limits, recordkeeping, and customer-support controls would PayPal require?"),
        ("9", "Would PayPal require a Canadian legal opinion, registration, corporation, business bank account, or other documentation before approval?"),
        ("10", "Which sandbox products, APIs, webhooks, partner credentials, and production-approval steps should the team use?"),
    ]
    add_table(doc, ["#", "Question"], questions, [700, 8660], header_fill=LIGHT_BLUE)

    add_heading(doc, "9. Requested next step", 1)
    add_text(doc, "Please connect us with a PayPal Platforms and Marketplaces representative or Canadian risk/compliance reviewer. We would like to provide any additional business details, product screens, flow diagrams, moderation policy drafts, or projected-volume assumptions required for a preliminary review. Please confirm the appropriate application path and provide a written determination or case reference before we build or activate live payment processing.")

    add_heading(doc, "Official PayPal references", 1)
    add_text(doc, "PayPal Platforms and Marketplaces overview: https://developer.paypal.com/platforms/overview", size=9.5, color=MUTED)
    add_text(doc, "PayPal Canada contact sales: https://securepayments.paypal.com/ca/business/contact-sales", size=9.5, color=MUTED)
    add_text(doc, "PayPal Canada Acceptable Use Policy: https://www.paypal.com/ca/legalhub/paypal/acceptableuse-full?locale.x=en_CA", size=9.5, color=MUTED)
    add_text(doc, "PayPal Canada delayed disbursement help: https://www.paypal.com/ca/cshelp/article/can-i-hold-money-before-disbursing-it-to-my-seller-with-paypal-multiparty-payments-ts2136", size=9.5, color=MUTED)
    add_text(doc, "This brief is a business and product summary for provider review, not legal advice and not evidence of PayPal approval.", before=10, after=0, size=9.5, color=MUTED, italic=True)

    path = DOCS / "paypal-platform-approval-brief.docx"
    doc.save(path)
    return path


def build_outreach():
    text = """fudnME — PayPal outreach package
Prepared: August 16, 2026

SUBMISSION ROUTE
PayPal Platforms and Marketplaces form:
https://developer.paypal.com/platforms/overview

PayPal Canada contact-sales form:
https://securepayments.paypal.com/ca/business/contact-sales

PayPal Canada sales phone: 1-866-357-0135

Use the Platforms and Marketplaces form first. If the form routes you to general sales, call the Canadian sales number and ask for a Platforms and Marketplaces / multiparty payments representative plus a Canadian risk/compliance review.

FORM VALUES
Company name: fudnME — pre-incorporation project
Website: https://fudnme.now
Country: Canada
Company type: Platform / marketplace / crowdfunding (select the closest available option)
Client annual volume: Pre-revenue; under CAD $100,000 planning range for the first 12 months
Existing PayPal Business account: No, if applicable

DESCRIPTION FIELD
We are developing fudnME, a Canadian consumer web platform where verified adults create public, lighthearted personal goals and other adults may give a small, unconditional monetary gift. A supporter may contribute a maximum of CAD $1 to a particular ask and may support up to ten asks per week. No product, service, reward, perk, equity, ownership, interest, loan, prize, raffle entry, contest entry, chance to win, or charitable tax receipt is provided. The target is aspirational. The product is pre-launch, no live funds are currently accepted, and the deployed prototype is a waitlist site at https://fudnme.now.

We are seeking preliminary written guidance on whether this model is permitted in Canada and whether PayPal Complete Payments / multiparty capabilities can support recipient onboarding, KYC/AML, partner fees, delayed disbursement, payouts, disputes, refunds, chargebacks, reserves, and reporting. The platform will not use a personal PayPal account or accept live funds before provider approval, incorporation, legal review, and required controls are complete. Please route this to Platforms and Marketplaces and Canadian risk/compliance, and provide a case or approval reference number.

FOLLOW-UP MESSAGE / EMAIL
Subject: Canadian pre-launch platform seeking PayPal multiparty approval review — fudnME

Hello PayPal Platforms and Marketplaces team,

I am building fudnME, a Canadian pre-launch web platform for small, unconditional peer-to-peer gifts. Verified adults will be able to post public personal goals; other adults may voluntarily give a maximum of CAD $1 to a particular ask, with a weekly limit on the number of asks a supporter can support.

This is not a sale of goods or services, a charity, an investment, a loan, a contest, a raffle, or a rewards-based crowdfunding campaign. Supporters receive nothing in exchange, and the target is aspirational. No live payments are currently accepted. Our deployed pre-launch prototype is https://fudnme.now.

We would like a preliminary written review of whether this model can be approved in Canada using PayPal Complete Payments / multiparty capabilities, including recipient onboarding, KYC/AML, partner fees, delayed disbursement, payout timing, disputes, refunds, chargebacks, reserves, and reporting. We also need guidance on the correct PayPal classification and whether a separate disclosed platform fee is permitted while the recipient receives the full stated CAD $1 gift amount.

We are currently pre-incorporation and pre-revenue. We will not open a production account under a founder’s personal name or process live funds until PayPal approval, Canadian legal review, and the required controls are in place.

Could you connect me with the appropriate Canadian Platforms and Marketplaces or risk/compliance representative? Please provide the correct application path, required documentation, and a case or approval reference number. I have attached a short approval brief with the proposed flow, controls, and specific questions.

Thank you,
[Founder legal name]
Founder, fudnME
[Business email]
[Phone number]
https://fudnme.now

PHONE SCRIPT
“Hi, I’m calling about a Canadian pre-incorporation consumer platform seeking a preliminary Platforms and Marketplaces review. The product is a pre-launch social micro-gifting platform: adult users may voluntarily give another adult a maximum of one Canadian dollar, with no product, service, reward, investment, contest, or charitable receipt. We need the correct multiparty/crowdfunding classification, onboarding path, and a written risk/compliance review before accepting live funds. Could you connect me with the Canadian Platforms and Marketplaces team and give me a case reference?”

DO NOT SAY
- “PayPal approved us.”
- “This is a charity.”
- “Users are investing.”
- “Supporters get a chance to win.”
- “We already process payments.”

ATTACHMENT
paypal-platform-approval-brief.docx
"""
    path = DOCS / "paypal-outreach-package.txt"
    path.write_text(text, encoding="utf-8")
    return path


if __name__ == "__main__":
    print(build_docx())
    print(build_outreach())
